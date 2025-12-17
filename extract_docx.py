#!/usr/bin/env python3
"""
Extract images and text from docx file
"""
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os
from pathlib import Path

def extract_images_from_docx(docx_path, output_dir="images"):
    """Extract all images from a docx file"""
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    doc = Document(docx_path)
    image_count = 0
    
    # Extract images from document.xml.rels
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_part = rel.target_part
            image_data = image_part.blob
            
            # Determine file extension
            content_type = image_part.content_type
            if "png" in content_type:
                ext = ".png"
            elif "jpeg" in content_type or "jpg" in content_type:
                ext = ".jpg"
            elif "gif" in content_type:
                ext = ".gif"
            else:
                ext = ".bin"
            
            filename = f"image_{image_count:03d}{ext}"
            filepath = os.path.join(output_dir, filename)
            
            with open(filepath, "wb") as f:
                f.write(image_data)
            
            print(f"Extracted: {filepath}")
            image_count += 1
    
    return image_count

def docx_to_text(docx_path, output_file="output.txt"):
    """Convert docx to plain text"""
    doc = Document(docx_path)
    
    with open(output_file, "w", encoding="utf-8") as f:
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                f.write(f"{para.text}\n")
    
    print(f"Text extracted to: {output_file}")

if __name__ == "__main__":
    docx_file = "Gemini_Suicide.docx"
    
    # Extract images
    print("Extracting images...")
    num_images = extract_images_from_docx(docx_file)
    print(f"Total images extracted: {num_images}\n")
    
    # Extract text
    print("Extracting text...")
    docx_to_text(docx_file, "output.txt")
